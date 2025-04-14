import streamlit as st
import docx
import requests
from bs4 import BeautifulSoup
import re
from urllib.parse import urlparse, urljoin, urlunparse
from docx.opc.constants import RELATIONSHIP_TYPE as RT
import fitz  # PyMuPDF
import io
import pandas as pd
from collections import OrderedDict, defaultdict
import asyncio
import aiohttp
from concurrent.futures import ThreadPoolExecutor
import nest_asyncio
from itertools import groupby
from operator import itemgetter
import mimetypes
import time
nest_asyncio.apply()

url_pattern = re.compile(
    r'\bhttps?://(?:www\.)?[-a-zA-Z0-9@:%._\+~#=]{1,256}\.[a-zA-Z]{2,6}\b(?:[-a-zA-Z0-9()@:%_\+.~#?&//=]*)')


def extract_hyperlinks_from_docx(docx_file):
    """Extract all hyperlinks from a DOCX file and use the same canonical-merging approach
       as the PDF extraction, so partial duplicates are merged into one final URL."""

    doc = docx.Document(docx_file)
    urls_list = []  # Keep the order of all URLs found
    url_counts = {}  # Count of each raw URL

    # 1) Collect URLs from hyperlink relationships
    for rel in doc.part.rels.values():
        if rel.reltype == RT.HYPERLINK:
            if rel._target.startswith('http'):
                urls_list.append(rel._target)
                url_counts[rel._target] = url_counts.get(rel._target, 0) + 1

    # 2) Collect URLs from paragraph text
    for paragraph in doc.paragraphs:
        found_urls = re.findall(url_pattern, paragraph.text)
        for url in found_urls:
            urls_list.append(url)
            url_counts[url] = url_counts.get(url, 0) + 1

    # 3) Collect URLs from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    found_urls = re.findall(url_pattern, paragraph.text)
                    for url in found_urls:
                        urls_list.append(url)
                        url_counts[url] = url_counts.get(url, 0) + 1

    # 4) Collect URLs from headers and footers
    for section in doc.sections:
        # Header
        if section.header:
            for paragraph in section.header.paragraphs:
                found_urls = re.findall(url_pattern, paragraph.text)
                for url in found_urls:
                    urls_list.append(url)
                    url_counts[url] = url_counts.get(url, 0) + 1
        # Footer
        if section.footer:
            for paragraph in section.footer.paragraphs:
                found_urls = re.findall(url_pattern, paragraph.text)
                for url in found_urls:
                    urls_list.append(url)
                    url_counts[url] = url_counts.get(url, 0) + 1

    # ------------------------------
    # Apply the same "canonical" merging approach as in the PDF logic
    # ------------------------------

    # Step 1: Determine a canonical URL for each raw URL
    #   For each URL, we look for longer URLs in our set that start with the same substring.
    #   The longest matching one is the canonical version.
    canonical_for = {}
    for url in url_counts:
        candidates = [cand for cand in url_counts if cand.startswith(url)]
        canonical_for[url] = max(candidates, key=len) if candidates else url

    # Step 2: Merge occurrence counts into each canonical URL
    merged_counts = {}
    for original_url, count in url_counts.items():
        canon = canonical_for[original_url]
        merged_counts[canon] = merged_counts.get(canon, 0) + count

    # Step 3: Build the final ordered list of canonical URLs
    final_urls = []
    seen_canonical = set()
    for url in urls_list:
        canon = canonical_for[url]
        if canon not in seen_canonical:
            final_urls.append(canon)
            seen_canonical.add(canon)

    # Step 4: Construct the final list of dictionaries with duplicate info
    url_info = []
    for url in final_urls:
        occurrence_count = merged_counts[url]
        url_info.append({
            'URL': url,
            'Occurrences': occurrence_count,
            'Is Duplicate': 'Yes' if occurrence_count > 1 else 'No'
        })

    return url_info


def extract_hyperlinks_from_pdf(file_bytes):
    """Extract all hyperlinks from a PDF file, merging duplicated links that may occur
    because of truncated URLs (i.e. when the text extraction cuts off a long URL)."""
    # Wrap the file bytes in a BytesIO stream for PyMuPDF
    pdf_stream = io.BytesIO(file_bytes.read())
    doc = fitz.open(stream=pdf_stream, filetype="pdf")

    urls_list = []     # This list tracks the order of appearance.
    url_counts = {}    # This dictionary tracks raw counts before merging.

    # Extract URLs from link annotations (these tend to be complete).
    for page in doc:
        links = page.get_links()
        for link in links:
            uri = link.get('uri')
            if uri and uri.startswith('http'):
                urls_list.append(uri)
                url_counts[uri] = url_counts.get(uri, 0) + 1

    # Extract URLs from the page text using regex.
    for page in doc:
        text = page.get_text()
        # Remove newline characters to reduce issues with line breaks splitting URLs.
        cleaned_text = text.replace('\n', '')
        found_urls = url_pattern.findall(cleaned_text)
        for url in found_urls:
            urls_list.append(url)
            url_counts[url] = url_counts.get(url, 0) + 1

    # === Step 1: Determine canonical URLs ===
    # For each unique URL, find if there's a longer URL in our collection that
    # starts with the same substring. That longer URL is assumed to be the correct (complete)
    # version.
    canonical_for = {}
    for url in url_counts:
        # Find all candidates in our keys that start with the current url.
        candidates = [cand for cand in url_counts if cand.startswith(url)]
        # If multiple candidates exist, choose the longest one; otherwise the url is canonical.
        canonical_for[url] = max(candidates, key=len) if candidates else url

    # === Step 2: Merge occurrence counts by canonical URL ===
    merged_counts = {}
    for original_url, count in url_counts.items():
        canon = canonical_for[original_url]
        merged_counts[canon] = merged_counts.get(canon, 0) + count

    # === Step 3: Produce final list in the order of first appearance (of the canonical URL) ===
    final_urls = []
    seen_canonical = set()
    for url in urls_list:
        canon = canonical_for[url]
        if canon not in seen_canonical:
            final_urls.append(canon)
            seen_canonical.add(canon)

    # Build a list with URL information.
    # Here, if the canonical URL was seen more than once, we mark it as a duplicate.
    url_info = []
    for url in final_urls:
        url_info.append({
            'URL': url,
            'Occurrences': merged_counts[url],
            'Is Duplicate': 'Yes' if merged_counts[url] > 1 else 'No'
        })

    return url_info


def get_domain(url):
    """Extract domain from URL."""
    try:
        parsed = urlparse(url)
        return parsed.netloc
    except:
        return url


def is_image_url(url):
    """Check if URL points to an image based on extension or content type."""
    parsed = urlparse(url)
    path = parsed.path.lower()
    return any(path.endswith(ext) for ext in ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.webp'])


async def check_url_async(session, url, delay=1, max_retries=3):
    """Asynchronously check if a URL is accessible and return its status and content."""
    retry_count = 0
    while retry_count < max_retries:
        try:
            headers = {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
            }

            # First request to check accessibility
            async with session.get(url, headers=headers, timeout=10, allow_redirects=True) as response:
                status = response.status
                final_url = str(response.url)

                # Handle redirects explicitly
                if final_url != url:
                    redirect_info = f"(Redirected to: {final_url})"
                else:
                    redirect_info = ""

                # Check if it's an image URL
                if is_image_url(final_url) or 'image' in response.headers.get('content-type', '').lower():
                    return {
                        'status': 'Success',
                        'status_code': status,
                        'content_preview': f"✓ Image content confirmed {redirect_info}"
                    }

                # Wait before getting content
                await asyncio.sleep(delay)

                # Second request to get content
                async with session.get(final_url, headers=headers, timeout=10) as content_response:
                    text = await content_response.text()

                    # Check for rate limiting indicators
                    rate_limit_indicators = [
                        'too many requests', 'rate limit', 'try again later', '429']
                    if any(indicator in text.lower() for indicator in rate_limit_indicators) or status == 429:
                        if retry_count < max_retries - 1:
                            retry_count += 1
                            # Wait 7 seconds before retrying
                            await asyncio.sleep(7)
                            continue

                    # Try to get readable content using BeautifulSoup
                    soup = BeautifulSoup(text, 'html.parser')

                    # Remove script and style elements
                    for script in soup(["script", "style"]):
                        script.decompose()

                    # Get text content
                    content = soup.get_text()

                    # Clean up text
                    lines = (line.strip() for line in content.splitlines())
                    chunks = (phrase.strip()
                              for line in lines for phrase in line.split("  "))
                    content = ' '.join(chunk for chunk in chunks if chunk)

                    # If content is just "Redirecting..." or similar, try to find more info
                    if content.lower().strip() in ['redirecting', 'redirecting...', '']:
                        # Try to find a redirect URL or title
                        redirect_url = soup.find(
                            'meta', attrs={'http-equiv': 'refresh'})
                        title = soup.find('title')
                        if redirect_url:
                            content = f"Redirect page. {redirect_info}"
                        elif title:
                            content = f"Page title: {title.text.strip()} {redirect_info}"
                        else:
                            content = f"Valid page with minimal content {redirect_info}"

                    return {
                        'status': 'Success',
                        'status_code': status,
                        'content_preview': content[:500] + '...' if len(content) > 500 else content
                    }

        except Exception as e:
            error_msg = str(e).lower()
            if any(msg in error_msg for msg in ['too many requests', 'rate limit', '429']):
                if retry_count < max_retries - 1:
                    retry_count += 1
                    await asyncio.sleep(7)  # Wait 7 seconds before retrying
                    continue
            return {
                'status': 'Error',
                'status_code': getattr(e, 'status', None) if hasattr(e, 'status') else None,
                'content_preview': str(e)
            }
        retry_count += 1

    return {
        'status': 'Error',
        'status_code': 429,
        'content_preview': 'Max retries exceeded - Rate limiting or server issues'
    }


async def process_domain_group(session, domain_urls, delay=1):
    """Process all URLs from the same domain with delays."""
    results = []
    for url, occurrences in domain_urls:
        # Wait before processing next URL from same domain
        await asyncio.sleep(delay)
        result = await check_url_async(session, url, delay)
        result['occurrences'] = occurrences
        results.append((url, result))
    return results


async def check_urls_batch(urls_dict):
    """Check multiple URLs concurrently, but handle same-domain URLs sequentially."""
    async with aiohttp.ClientSession() as session:
        results = {'success': [], 'failed': []}

        # Group URLs by domain
        domain_groups = defaultdict(list)
        for url, occurrences in urls_dict.items():
            domain = get_domain(url)
            domain_groups[domain].append((url, occurrences))

        # Create tasks for each domain group
        tasks = []
        for domain, domain_urls in domain_groups.items():
            task = asyncio.ensure_future(
                process_domain_group(session, domain_urls))
            tasks.append(task)

        # Process domain groups concurrently (but URLs within each domain sequentially)
        total_urls = len(urls_dict)
        processed_urls = 0

        # Wait for all domain groups to complete
        for domain_results in asyncio.as_completed(tasks):
            domain_processed = await domain_results
            for url, result in domain_processed:
                if result['status'] == 'Success':
                    results['success'].append({'url': url, **result})
                else:
                    results['failed'].append({'url': url, **result})

                processed_urls += 1
                # Update progress
                if 'progress_bar' in st.session_state:
                    st.session_state.progress_bar.progress(
                        processed_urls / total_urls)
                if 'status_text' in st.session_state:
                    st.session_state.status_text.text(
                        f"Processed {processed_urls} of {total_urls} URLs")

        return results


def main():
    st.set_page_config(page_title="Document URL Checker", layout="wide")

    # Add custom CSS
    st.markdown("""
        <style>
        .success-url { color: #28a745; }
        .failed-url { color: #dc3545; }
        .stProgress > div > div > div > div { background-color: #28a745; }
        .duplicate-url { background-color: #e8eaed; }
        </style>
    """, unsafe_allow_html=True)

    st.title("📄 Document URL Checker")
    st.write("Upload a Word document (.docx) or a PDF file (.pdf) to check all URLs contained within it.")

    # File uploader
    uploaded_file = st.file_uploader(
        "Choose a DOCX or PDF file", type=["docx", "pdf"])

    if uploaded_file is not None:
        try:
            # Create a bytes buffer for the uploaded file
            file_bytes = io.BytesIO(uploaded_file.read())
            file_extension = uploaded_file.name.lower().split('.')[-1]

            with st.spinner('🔍 Extracting URLs from document...'):
                if file_extension == "docx":
                    url_info = extract_hyperlinks_from_docx(file_bytes)
                elif file_extension == "pdf":
                    url_info = extract_hyperlinks_from_pdf(file_bytes)
                else:
                    st.warning(
                        "❌ Unsupported file type. Please upload a .docx or .pdf file.")
                    st.stop()

            if not url_info:
                st.warning("⚠️ No URLs found in the document.")
            else:
                total_urls = len(url_info)
                unique_urls = len(set(item['URL'] for item in url_info))
                duplicate_count = total_urls - unique_urls

                st.success(
                    f"✨ Found {total_urls} total URLs ({unique_urls} unique, {duplicate_count} duplicates)")

            # Display URLs in a table
            url_df = pd.DataFrame(url_info)

            # Style the dataframe to highlight duplicates
            def highlight_duplicates(row):
                return ['background-color: #e8eaed' if row['Is Duplicate'] == 'Yes' else '' for _ in row]

            styled_df = url_df.style.apply(highlight_duplicates, axis=1)
            st.dataframe(styled_df, use_container_width=True)

            # Add a button to start checking URLs
            if st.button("🚀 Start Checking URLs"):
                # Initialize progress bar and status
                st.session_state.progress_bar = st.progress(0)
                st.session_state.status_text = st.empty()

                # Create two columns for results
                col1, col2 = st.columns(2)

                # Check URLs concurrently
                unique_urls = {item['URL']: item['Occurrences']
                               for item in url_info}
                results = asyncio.run(check_urls_batch(unique_urls))

                # Clear the status text
                st.session_state.status_text.empty()

                # Show results in columns
                with col1:
                    st.markdown("### ✅ Working URLs")
                    if results['success']:
                        success_data = []
                        for result in results['success']:
                            success_data.append({
                                'URL': result['url'],
                                'Status': result['status_code'],
                                'Occurrences': result['occurrences'],
                                'Content Preview': result['content_preview'][:150] + '...' if len(result['content_preview']) > 150 else result['content_preview']
                            })
                        success_df = pd.DataFrame(success_data)
                        st.dataframe(success_df, use_container_width=True)
                    else:
                        st.write("No working URLs found.")

                with col2:
                    st.markdown("### ❌ Failed URLs")
                    if results['failed']:
                        failed_data = []
                        for result in results['failed']:
                            failed_data.append({
                                'URL': result['url'],
                                'Status': result['status_code'],
                                'Occurrences': result['occurrences'],
                                'Error Message': result['content_preview'][:150] + '...' if len(result['content_preview']) > 150 else result['content_preview']
                            })
                        failed_df = pd.DataFrame(failed_data)
                        st.dataframe(failed_df, use_container_width=True)
                    else:
                        st.write("No failed URLs found.")

                # ✅ Export functionality with working CSV download
                if results['success'] or results['failed']:
                    export_data = []
                    for result in results['success'] + results['failed']:
                        export_data.append({
                            'URL': result['url'],
                            'Status': result.get('status', 'N/A'),
                            'Status Code': result.get('status_code', 'N/A'),
                            'Times Found': result['occurrences'],
                            'Details': result['content_preview'][:200]
                        })
                    df = pd.DataFrame(export_data)
                    csv = df.to_csv(index=False).encode('utf-8')

                    st.download_button(
                        label="📥 Download CSV",
                        data=csv,
                        file_name="url_check_results.csv",
                        mime="text/csv"
                    )

        except Exception as e:
            st.error(f"An error occurred while processing the file: {str(e)}")
            st.exception(e)


if __name__ == "__main__":
    main()
