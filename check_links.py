import docx
import requests
from bs4 import BeautifulSoup
import re
from urllib.parse import urlparse
from docx.opc.constants import RELATIONSHIP_TYPE as RT

def extract_hyperlinks_from_docx(docx_path):
    """Extract all hyperlinks from a DOCX file, including embedded ones."""
    doc = docx.Document(docx_path)
    urls = set()  # Using a set to avoid duplicates
    
    # Get all hyperlinks through document relationships
    # This is the most reliable method for getting hyperlinks
    for rel in doc.part.rels.values():
        if rel.reltype == RT.HYPERLINK:
            # Some hyperlinks might be relative or internal, we want only URLs
            if rel._target.startswith('http'):
                urls.add(rel._target)
    
    # Also check for any URLs in the text (as backup)
    url_pattern = r'http[s]?://(?:[a-zA-Z]|[0-9]|[$-_@.&+]|[!*\\(\\),]|(?:%[0-9a-fA-F][0-9a-fA-F]))+'
    
    # Check paragraphs
    for paragraph in doc.paragraphs:
        found_urls = re.findall(url_pattern, paragraph.text)
        urls.update(found_urls)
    
    # Check tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    found_urls = re.findall(url_pattern, paragraph.text)
                    urls.update(found_urls)
    
    # Get hyperlinks from headers and footers
    for section in doc.sections:
        # Header
        if section.header:
            for paragraph in section.header.paragraphs:
                found_urls = re.findall(url_pattern, paragraph.text)
                urls.update(found_urls)
        # Footer
        if section.footer:
            for paragraph in section.footer.paragraphs:
                found_urls = re.findall(url_pattern, paragraph.text)
                urls.update(found_urls)
    
    return sorted(list(urls))  # Convert back to sorted list

def check_url(url):
    """Check if a URL is accessible and return its status and content."""
    try:
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
        }
        response = requests.get(url, headers=headers, timeout=10)
        response.raise_for_status()
        
        # Try to get readable content using BeautifulSoup
        soup = BeautifulSoup(response.text, 'html.parser')
        
        # Remove script and style elements
        for script in soup(["script", "style"]):
            script.decompose()
            
        # Get text content
        text = soup.get_text()
        
        # Clean up text
        lines = (line.strip() for line in text.splitlines())
        chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
        text = ' '.join(chunk for chunk in chunks if chunk)
        
        return {
            'status': 'Success',
            'status_code': response.status_code,
            'content_preview': text[:500] + '...' if len(text) > 500 else text
        }
    except requests.RequestException as e:
        return {
            'status': 'Error',
            'status_code': getattr(e.response, 'status_code', None) if hasattr(e, 'response') else None,
            'content_preview': str(e)
        }

def main():
    docx_path = "INIT12_Type 1 Report 2024 DRAFT.docx"
    print(f"\nExtracting URLs from {docx_path}...")
    
    try:
        urls = extract_hyperlinks_from_docx(docx_path)
        print(f"\nFound {len(urls)} unique URLs in the document:")
        for i, url in enumerate(urls, 1):
            print(f"\n{i}. {url}")
        
        if not urls:
            print("\nNo URLs found in the document. This might indicate an issue with the document format or the way hyperlinks are stored.")
            return
            
        print("\nChecking each URL...")
        print("-" * 80)
        
        results = {'success': [], 'failed': []}
        for i, url in enumerate(urls, 1):
            print(f"\nChecking URL {i}/{len(urls)}: {url}")
            result = check_url(url)
            if result['status'] == 'Success':
                results['success'].append({'url': url, **result})
                print(f"✅ Success (Status: {result['status_code']})")
            else:
                results['failed'].append({'url': url, **result})
                print(f"❌ Failed: {result['content_preview']}")
        
        # Print summary
        print("\n" + "=" * 80)
        print("SUMMARY")
        print("=" * 80)
        
        print("\n✅ Working URLs:")
        for result in results['success']:
            print(f"\n{result['url']} (Status: {result['status_code']})")
            print(f"Preview: {result['content_preview'][:200]}...")
        
        print("\n❌ Failed URLs:")
        for result in results['failed']:
            print(f"\n{result['url']} (Status: {result['status_code']})")
            print(f"Error: {result['content_preview']}")
        
        print("\nSTATISTICS:")
        print(f"Total URLs: {len(urls)}")
        print(f"Working URLs: {len(results['success'])}")
        print(f"Failed URLs: {len(results['failed'])}")
        
    except Exception as e:
        print(f"An error occurred while processing the file: {str(e)}")
        import traceback
        print("\nFull error traceback:")
        print(traceback.format_exc())

if __name__ == "__main__":
    main() 