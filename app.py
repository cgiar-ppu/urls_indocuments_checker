import streamlit as st
import docx
import requests
from bs4 import BeautifulSoup
import re
from urllib.parse import urlparse
import io

def extract_urls_from_docx(docx_file):
    """Extract all URLs from a DOCX file."""
    doc = docx.Document(docx_file)
    urls = []
    
    # Regular expression for URL matching
    url_pattern = r'http[s]?://(?:[a-zA-Z]|[0-9]|[$-_@.&+]|[!*\\(\\),]|(?:%[0-9a-fA-F][0-9a-fA-F]))+'
    
    # Extract URLs from paragraphs
    for paragraph in doc.paragraphs:
        found_urls = re.findall(url_pattern, paragraph.text)
        urls.extend(found_urls)
    
    # Extract URLs from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                found_urls = re.findall(url_pattern, cell.text)
                urls.extend(found_urls)
    
    return list(set(urls))  # Remove duplicates

def check_url(url):
    """Check if a URL is accessible and return its status and content."""
    try:
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
        }
        response = requests.get(url, headers=headers, timeout=10)
        response.raise_for_status()
        
        # Try to get readable content using BeautifulSoup
        soup = BeautifulSoup(response.text, 'html.parser')
        
        # Remove script and style elements
        for script in soup(["script", "style"]):
            script.decompose()
            
        # Get text content
        text = soup.get_text()
        
        # Clean up text
        lines = (line.strip() for line in text.splitlines())
        chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
        text = ' '.join(chunk for chunk in chunks if chunk)
        
        return {
            'status': 'Success',
            'status_code': response.status_code,
            'content_preview': text[:500] + '...' if len(text) > 500 else text
        }
    except requests.RequestException as e:
        return {
            'status': 'Error',
            'status_code': getattr(e.response, 'status_code', None) if hasattr(e, 'response') else None,
            'content_preview': str(e)
        }

def main():
    st.set_page_config(page_title="Document URL Checker", layout="wide")
    
    st.title("Document URL Checker")
    st.write("Upload a .docx file to check all URLs contained within it.")
    
    uploaded_file = st.file_uploader("Choose a DOCX file", type="docx")
    
    if uploaded_file is not None:
        try:
            # Create a bytes buffer for the uploaded file
            docx_bytes = io.BytesIO(uploaded_file.read())
            
            with st.spinner('Extracting URLs from document...'):
                urls = extract_urls_from_docx(docx_bytes)
            
            if not urls:
                st.warning("No URLs found in the document.")
            else:
                st.success(f"Found {len(urls)} unique URLs in the document.")
                
                # Create a progress bar
                progress_bar = st.progress(0)
                
                # Check each URL
                results = []
                for i, url in enumerate(urls):
                    with st.spinner(f'Checking URL {i+1} of {len(urls)}: {url}'):
                        result = check_url(url)
                        results.append({'url': url, **result})
                        progress_bar.progress((i + 1) / len(urls))
                
                # Display results
                st.subheader("Results")
                
                # Success URLs
                st.write("### Working URLs")
                success_urls = [r for r in results if r['status'] == 'Success']
                if success_urls:
                    for result in success_urls:
                        with st.expander(f"✅ {result['url']} (Status: {result['status_code']})"):
                            st.text_area("Content Preview:", result['content_preview'], height=150)
                else:
                    st.write("No working URLs found.")
                
                # Failed URLs
                st.write("### Failed URLs")
                failed_urls = [r for r in results if r['status'] == 'Error']
                if failed_urls:
                    for result in failed_urls:
                        with st.expander(f"❌ {result['url']} (Status: {result['status_code']})"):
                            st.error(result['content_preview'])
                else:
                    st.write("No failed URLs found.")
                
        except Exception as e:
            st.error(f"An error occurred while processing the file: {str(e)}")

if __name__ == "__main__":
    main() 